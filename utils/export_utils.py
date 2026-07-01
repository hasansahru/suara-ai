from __future__ import annotations
import io, json, csv

def generate_txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")

def generate_json_bytes(data: dict, indent: int = 2) -> bytes:
    return json.dumps(data, ensure_ascii=False, indent=indent).encode("utf-8")

def generate_csv_bytes(headers: list, rows: list) -> bytes:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(headers)
    writer.writerows(rows)
    return output.getvalue().encode("utf-8")

def generate_docx_bytes(markdown_text: str, title: str = "Suara AI Output") -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise ImportError("pip install python-docx")
    if not markdown_text:
        return b""
    import re
    doc = Document()
    doc.add_heading(title, level=0)
    for line in markdown_text.split("\n"):
        s = line.strip()
        if not s:
            doc.add_paragraph("")
        elif s.startswith("#### "):
            doc.add_heading(s[5:], 4)
        elif s.startswith("### "):
            doc.add_heading(s[4:], 3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], 2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], 1)
        elif s in ("---", "***", "___"):
            doc.add_paragraph("─" * 40)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", s), style="List Number")
        else:
            doc.add_paragraph(re.sub(r"\*\*(.+?)\*\*", r"\1", s))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()